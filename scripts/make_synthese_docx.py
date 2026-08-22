#!/usr/bin/env python
"""Génère MS-VICRegL_synthese.docx — synthèse problématiques + résultats (python-docx)."""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

BLUE = RGBColor(0x29, 0x52, 0xCC)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x1E, 0x88, 0x49)
ORANGE = RGBColor(0xB9, 0x77, 0x0E)
GREY = RGBColor(0x5D, 0x6D, 0x7E)
DARK = RGBColor(0x1C, 0x28, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ---- base style ----
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.font.color.rgb = DARK

for name, sz, col in (("Heading 1", 15, DARK), ("Heading 2", 12, BLUE)):
    h = doc.styles[name]
    h.font.name = "Calibri"
    h.font.size = Pt(sz)
    h.font.bold = True
    h.font.color.rgb = col

sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.75)
sec.left_margin = sec.right_margin = Inches(0.83)


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)


def run(p, text, *, bold=False, italic=False, color=None, size=None, mono=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    if mono:
        r.font.name = "Consolas"
    return r


def para(runs, *, align=None, after=6, line=1.15):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.line_spacing = line
    for spec in runs:
        run(p, spec[0], **spec[1])
    return p


def bullet(runs):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    for spec in runs:
        run(p, spec[0], **spec[1])
    return p


def r(text, **kw):
    return (text, kw)


def add_table(header, rows, widths):
    """rows: list of (cells, colors_list_or_None, bold_bool)."""
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    # header
    for i, (txt, w) in enumerate(zip(header, widths)):
        c = tbl.rows[0].cells[i]
        c.width = Inches(w)
        shade(c, "1C2833")
        pc = c.paragraphs[0]
        pc.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        for j, line in enumerate(txt.split("\n")):
            if j:
                pc = c.add_paragraph()
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(pc, line, bold=True, color=WHITE, size=9.5)
    # body
    for ridx, (cells, colors, bold) in enumerate(rows):
        row = tbl.add_row()
        for i, txt in enumerate(cells):
            c = row.cells[i]
            c.width = Inches(widths[i])
            shade(c, "F4F6F7" if ridx % 2 else "FFFFFF")
            pc = c.paragraphs[0]
            pc.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            pc.paragraph_format.space_after = Pt(2)
            col = colors[i] if colors else DARK
            for j, line in enumerate(str(txt).split("\n")):
                if j:
                    pc = c.add_paragraph()
                    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run(pc, line, bold=bold, color=col, size=9.5)
    return tbl


def hr(p, top=False):
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    edge = OxmlElement("w:top" if top else "w:bottom")
    edge.set(qn("w:val"), "single"); edge.set(qn("w:sz"), "12" if not top else "8")
    edge.set(qn("w:space"), "6"); edge.set(qn("w:color"), "2952CC" if not top else "D5D8DC")
    bd.append(edge); pPr.append(bd)


# ===================== TITRE =====================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(t, "Identification bactérienne par MALDI-TOF", bold=True, size=17, color=DARK)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(s, "Un CNN 1D auto-supervisé (VICRegL) invariant aux artefacts de centre", size=13, color=BLUE)
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(s2, "Synthèse des problématiques et des résultats", italic=True, size=11, color=GREY)
s3 = doc.add_paragraph(); s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(s3, "Jeu DRIAMS (4 centres hospitaliers) · comparaison au pipeline classique MSclassifR",
    size=9.5, color=GREY)
hr(s3)

# ===================== 1. CONTEXTE =====================
doc.add_heading("1. Contexte et problématique", level=1)
para([r("L’identification bactérienne par spectrométrie MALDI-TOF est routinière, mais les "
        "modèles souffrent d’un "), r("effet de centre", bold=True),
      r(" : calibration m/z, ligne de base, gain du détecteur et matrice varient d’un hôpital "
        "et d’un appareil à l’autre. Un classifieur entraîné dans un centre "),
      r("se dégrade sur un autre", italic=True),
      r(", ce qui limite le déploiement multi-sites.")], align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para([r("Le pipeline de référence (type "), r("MSclassifR", bold=True),
      r(") neutralise ces artefacts par un pré-traitement lourd fait main (stabilisation de "
        "variance, ondelettes, SNIP, normalisation TIC, binning 3 Da, correction de lot). "),
      r("Question de ce travail : ", bold=True, color=BLUE),
      r("peut-on plutôt "), r("apprendre l’invariance aux artefacts", bold=True),
      r(", à partir de spectres quasi-bruts, sans ce pré-traitement expert ?")],
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ===================== 2. APPROCHE =====================
doc.add_heading("2. Approche proposée", level=1)
para([r("Un encodeur "), r("CNN 1D (ResNet-1D)", bold=True),
      r(" est pré-entraîné en "), r("auto-supervision VICRegL", bold=True),
      r(" — sans étiquettes. Les augmentations "), r("simulent les artefacts MALDI", italic=True),
      r(" (déformation de calibration m/z, ligne de base, gain, bruit, perte de pics) : le "
        "réseau apprend à ignorer ces nuisances tout en préservant la structure discriminante "
        "des pics (biomarqueurs) via le critère local de VICRegL.")],
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para([r("Entrée : ", bold=True),
      r("spectre brut ré-échantillonné (2000–20000 Da, 6000 points) + normalisation TIC "
        "seulement.  "),
      r("Évaluation : ", bold=True),
      r("sonde linéaire sur features gelées, comparée à un Random Forest sur "),
      r("binned_6000", mono=True, size=10),
      r(" (représentation DRIAMS standard, pré-traitement complet).")],
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para([r("Données DRIAMS : ", bold=True),
      r("4 centres — B (Bâle-Land), C (Aarau), D (Viollier), A (Bâle). Étiquettes "),
      r("Biotyper (MALDI)", bold=True, color=ORANGE),
      r(". Métrique principale : "), r("balanced accuracy", italic=True),
      r(" (classes déséquilibrées).")], align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ===================== 3. RÉSULTATS =====================
doc.add_heading("3. Résultats — vue d’ensemble", level=1)
add_table(
    ["Expérience", "Résultat clé", "Verdict"],
    [
        (["ID espèces cross-centre B↔C",
          "VICRegL 0,95–0,99 (stable) vs RF 0,68–1,00 (erratique) — ~8× plus stable",
          "Favorable*"], [DARK, DARK, GREEN], False),
        (["Espèces proches (Enterobacter, Strepto, Klebsiella)",
          "VICRegL > RF sur les 3 groupes (McNemar p<0,01)", "Favorable*"],
         [DARK, DARK, GREEN], False),
        (["Résistance antibiotique (AMR)",
          "RF bat VICRegL, aucun gain de stabilité", "Défavorable"],
         [DARK, DARK, RED], False),
        (["Hold-out centre D (jamais vu)",
          "RF 0,935 > VICRegL 0,886 — renversement", "Défavorable"],
         [DARK, DARK, RED], False),
        (["D ajouté au pré-entraînement (non-labellisé)",
          "0,886 → 0,915 : écart au RF divisé par 2", "Partiel"],
         [DARK, DARK, ORANGE], False),
    ],
    [2.35, 3.25, 1.15],
)
para([r("* dans le cadre où les deux centres ont été vus (non-labellisés) au pré-entraînement "
        "— voir §4.", italic=True, size=9, color=GREY)], after=8)

# ===================== 4. POINT CRITIQUE =====================
doc.add_heading("4. Le point critique : adaptation ≠ généralisation", level=1)
para([r("Le résultat marquant « ~8× plus stable » (B↔C) est obtenu quand "),
      r("les deux centres ont été vus, non-labellisés, au pré-entraînement", bold=True),
      r(" : c’est une "), r("adaptation de domaine à cible disponible", italic=True),
      r(", pas de la généralisation à un centre neuf. Le "), r("test strict", bold=True),
      r(" — centre D tenu totalement à l’écart du SSL — "),
      r("inverse la conclusion", bold=True, color=RED),
      r(" : le pipeline classique l’emporte.")], align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_table(
    ["Condition (éval. sur D)", "Encodeur B∪C\n(D jamais vu)",
     "Encodeur B∪C∪D\n(D vu non-labellisé)", "RF classique"],
    [
        (["(B+C) → D  (centre inédit)", "0,886", "0,915", "0,935"],
         [DARK, RED, ORANGE, GREEN], True),
        (["Plafond D → D (sonde sur D)", "0,914", "0,931", "—"], None, False),
        (["B+C en intra-domaine (plafond)", "0,988", "0,990", "0,889"], None, False),
        (["AUROC « quel centre ? » sur features", "0,997", "0,999", "—"],
         [DARK, RED, RED, DARK], False),
    ],
    [2.55, 1.55, 1.75, 1.15],
)
para([r("Balanced accuracy ; 10 espèces communes B∩C∩D ; test = 8797 spectres de D.",
        italic=True, size=9, color=GREY)], after=8)
para([r("Diagnostic : ", bold=True, color=BLUE),
      r("l’AUROC « quel centre ? » reste ~1,0 même après avoir vu D — les features "),
      r("encodent encore parfaitement le centre", bold=True),
      r(". L’invariance revendiquée "), r("n’est pas atteinte", bold=True, color=RED),
      r(" ; ce qui fonctionnait en B↔C relevait de l’adaptation. Voir D en SSL enrichit la "
        "représentation (gain réel, sans labels) mais "),
      r("ne restaure ni la parité avec le RF ni l’invariance", italic=True), r(".")],
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ===================== 5. FORCES / LIMITES =====================
doc.add_heading("5. Forces et limites", level=1)
doc.add_heading("Forces", level=2)
bullet([r("Pas de pré-traitement expert", bold=True),
        r(" : spectre quasi-brut + TIC ; plafond intra-domaine élevé (0,99).")])
bullet([r("Extension d’un centre sans aucun label", bold=True),
        r(" : l’ajouter au SSL (non-supervisé) réduit de moitié l’écart au pipeline classique.")])
bullet([r("Diagnostic rigoureux", bold=True),
        r(" : décomposition plafond/transfert, classifieur de domaine, McNemar apparié, "
          "IC bootstrap.")])
doc.add_heading("Limites", level=2)
bullet([r("Circularité des étiquettes", bold=True, color=ORANGE),
        r(" : DRIAMS est labellisé Biotyper (MALDI), pas NGS — 0 Shigella dans B+C. On ne peut "
          "pas prétendre « battre le MALDI » avec des labels MALDI.")])
bullet([r("Invariance non atteinte", bold=True),
        r(" : AUROC domaine ~1,0 ; échec sur un centre vraiment inédit.")])
bullet([r("AMR", bold=True),
        r(" : les augmentations (perte de pics) effacent le signal faible que la résistance "
          "requiert.")])

# ===================== 6. CONCLUSION =====================
doc.add_heading("6. Conclusion et perspectives", level=1)
para([r("Le pipeline est un "), r("prototype de méthode prometteur", bold=True),
      r(", pas un système supérieur au classique : il excelle en intra-domaine et en "
        "adaptation, mais "), r("l’invariance apprise ne généralise pas encore", italic=True),
      r(" à un centre inédit, où le pré-traitement fait main reste devant.")],
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para([r("Deux leviers pour la suite :", bold=True, color=BLUE)])
bullet([r("Terme d’invariance de centre explicite", bold=True),
        r(" dans la loss (domaine-adversarial / CORAL) — attaque directement l’AUROC ~1,0, "
          "cause racine de l’échec.")])
bullet([r("Données à étiquettes moléculaires (WGS)", bold=True),
        r(" — seul moyen de lever la circularité et de tester la vraie hypothèse : l’invariance "
          "capte-t-elle un signal que les labels MALDI ratent ?")])
para([r("Ce dernier point ouvre une "), r("collaboration naturelle", bold=True, color=GREEN),
      r(" avec les travaux sur données MALDI appariées à des labels moléculaires "
        "(identification fine de sous-espèces, rejet par confiance).")],
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
foot = doc.add_paragraph()
foot.paragraph_format.space_before = Pt(14)
hr(foot, top=True)
run(foot, "Reproductibilité : encodeur ResNet-1D, PyTorch MPS (Apple M3 Pro). Résultats "
          "régénérables via scripts/08_holdout_D.py et 09_ssl_BCD.py.",
    italic=True, size=8.5, color=GREY)

out = Path(__file__).resolve().parents[1] / "MS-VICRegL_synthese.docx"
doc.save(out)
print("OK ->", out)
